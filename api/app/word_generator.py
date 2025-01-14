from docx import Document
import tempfile
import os

async def create_word(project_name: str, summary_text: str) -> str:
    """議事録を作成して一時ファイルパスを返す関数"""
    # ファイル名を指定
    file_name = f"議事録_{project_name}.docx"
    # 一時ディレクトリを取得し、完全なパスを作成
    temp_dir = tempfile.gettempdir()
    temp_path = os.path.join(temp_dir, file_name)

    # Wordファイル作成
    document = Document()
    document.add_heading("議事録", level=1)
    document.add_paragraph(f"プロジェクト名: {project_name}")
    document.add_paragraph(f"要約: {summary_text}")
    document.save(temp_path)

    print(f"作成されたWordファイル: {temp_path}")
    return temp_path


async def cleanup_file(file_path: str):
    """一時ファイルを削除"""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"削除成功: {file_path}")
    except Exception as e:
        print(f"削除失敗: {e}")
